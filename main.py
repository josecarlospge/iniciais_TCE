#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Sistema de Extração de Certidões de Débito do TCE/PI
Extrai dados de certidões PDF e gera petições iniciais de execução
"""

import re
import json
import sqlite3
import os
import subprocess
import platform
from datetime import datetime
from pathlib import Path
import pdfplumber
from pypdf import PdfReader, PdfWriter
from enviador_peticao import abrir_enviador


# ============================================================
# EXTRAÇÃO DE DADOS DA CERTIDÃO
# ============================================================

def extrair_texto_pdf(caminho_pdf: str) -> str:
    """
    Extrai texto completo de um PDF.
    Estratégia 1: pdfplumber (rápido, funciona na maioria dos PDFs).
    Estratégia 2: OCR via pdf2image + pytesseract (fallback para PDFs com
                  fontes de encoding privado, como os gerados pelo TCE/PI
                  com assinatura digital via PScript5/Acrobat Distiller).
    """
    texto = ""
    with pdfplumber.open(caminho_pdf) as pdf:
        for pagina in pdf.pages:
            t = pagina.extract_text()
            if t:
                texto += t + "\n"

    # Detecta encoding quebrado: texto cheio de "(cid:XX)" ou "/iXXX"
    # indica que a fonte usa mapeamento privado — aciona OCR
    cid_count = texto.count("(cid:")
    total_chars = len(texto.strip())
    encoding_quebrado = (cid_count > 5) or (total_chars < 50)

    if encoding_quebrado:
        print("  ⚠️  Encoding de fonte não mapeável — usando OCR...")
        texto = _extrair_texto_ocr(caminho_pdf)

    return texto


def _extrair_texto_ocr(caminho_pdf: str) -> str:
    """Extrai texto via OCR (pdf2image + pytesseract). Requer poppler e tesseract."""
    from pdf2image import convert_from_path
    import pytesseract

    # Usa a mesma configuracao centralizada do gestor_enderecos
    try:
        from gestor_enderecos import _configurar_ocr
        poppler_path, _ = _configurar_ocr()
    except ImportError:
        # Fallback manual caso gestor_enderecos nao esteja disponivel
        from pathlib import Path
        base_dir = Path(__file__).parent
        poppler_local = base_dir / "poppler-25.12.0" / "Library" / "bin"
        poppler_path = str(poppler_local) if poppler_local.exists() else None
        tess_win = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        pytesseract.pytesseract.tesseract_cmd = tess_win
        os.environ["TESSDATA_PREFIX"] = r"C:\Program Files\Tesseract-OCR\tessdata"

    kwargs = {"dpi": 200}
    if poppler_path:
        kwargs["poppler_path"] = poppler_path

    imgs = convert_from_path(caminho_pdf, **kwargs)

    partes = []
    for img in imgs:
        texto = pytesseract.image_to_string(img, lang="por+eng")
        partes.append(texto)

    return "\n".join(partes)


def _normalizar_acordao(s: str) -> str:
    """
    Normaliza número de acórdão:
      '185/2023 - SPL'   → '185/2023'
      '185-B/2023 - SPL' → '185-B/2023'   (preserva o sufixo antes da barra)
      '90/2024-SPL'      → '90/2024'
    Remove apenas os sufixos após o ANO (ex: ' - SPL'), mantendo prefixos
    alfanuméricos que fazem parte do número (ex: -B, -C).
    """
    s = s.strip()
    # Remove sufixos após o ano: "/2023 - SPL" → "/2023", "/2024-SPL" → "/2024"
    s = re.sub(r'(/\d{4})\s*[-–]\s*\w+', r'\1', s)
    return s


def extrair_dados_certidao(texto: str) -> dict:
    """
    Extrai os dados estruturados de uma certidão de débito do TCE/PI.

    Estratégia de extração de responsáveis:
      - Fonte primária: parágrafo "CERTIFICO, ainda..." que lista os
        responsáveis solidários REMANESCENTES (já descontadas as exclusões
        por recurso). É o parágrafo que define quem efetivamente deve.
      - Fonte secundária (fallback): primeiro parágrafo "CERTIFICO, para os
        fins...", caso o segundo não exista ou não contenha CPF/CNPJ.
      - O acórdão de cada responsável é buscado no primeiro parágrafo,
        onde ficam os acórdãos individuais (185/2023, 185-B/2023, etc.).
    """
    texto_flat = re.sub(r'\s+', ' ', texto)

    dados = {
        "numero_processo": None,
        "valor_atualizado": None,
        "data_atualizacao": None,
        "acordao_origem": None,
        "responsaveis": []
    }

    # -------------------------------------------------------
    # 1) Número do processo (primeira ocorrência)
    # -------------------------------------------------------
    m = re.search(
        r'processo\s+n[oº°]?\s*(TC[/\s]\d{6}[/\s]\d{4})',
        texto_flat, re.IGNORECASE
    )
    if m:
        dados["numero_processo"] = re.sub(r'\s', '/', m.group(1))

    # -------------------------------------------------------
    # 2) Valor atualizado (no parágrafo "restituir ... R$ X")
    # -------------------------------------------------------
    m = re.search(
        r'(?:restituir[^R]{0,80}|[Vv]alor\s+(?:final|do\s+débito\s+atualizado)[^R]{0,20})'
        r'R\$\s*([\d.,]+)',
        texto_flat, re.IGNORECASE
    )
    if m:
        dados["valor_atualizado"] = "R$ " + m.group(1)
    else:
        # Fallback: último valor monetário no texto
        todos = re.findall(r'R\$\s*(\d{1,3}(?:\.\d{3})*,\d{2})', texto_flat)
        if todos:
            dados["valor_atualizado"] = "R$ " + todos[-1]

    # -------------------------------------------------------
    # 3) Data de atualização
    # -------------------------------------------------------
    for pat in [
        r'[Aa]tualizado\s+em[:\s]+(\d{2}/\d{2}/\d{4})',
        r'até\s+(\d{2}/\d{2}/\d{4})',
    ]:
        m = re.search(pat, texto_flat)
        if m:
            dados["data_atualizacao"] = m.group(1)
            break

    # -------------------------------------------------------
    # 4) Acórdão de origem (primeiro acórdão do texto inteiro)
    # -------------------------------------------------------
    m = re.search(
        r'[Aa]córd[ãa]o\s+n[oº°]?\s*([\d]+(?:-[A-Z]+)?/\d{4})',
        texto_flat
    )
    if m:
        dados["acordao_origem"] = _normalizar_acordao(m.group(1))

    # -------------------------------------------------------
    # 5) Responsáveis — fonte: parágrafo "CERTIFICO, ainda"
    #    (responsáveis solidários remanescentes = quem efetivamente deve)
    # -------------------------------------------------------
    trecho_resp = texto_flat   # default: texto inteiro

    m_segundo = re.search(
        r'CERTIFICO,\s+ainda[^.]*\.',   # até o primeiro ponto final
        texto_flat, re.IGNORECASE
    )
    # "CERTIFICO, ainda" pode se estender por várias frases; pegamos até
    # o fim do parágrafo (fim da sentença com "substituí-la" ou próximo ponto
    # seguido de nova seção)
    m_segundo = re.search(
        r'(CERTIFICO,\s+ainda.+?substituí-la\.)',
        texto_flat, re.IGNORECASE | re.DOTALL
    )
    if not m_segundo:
        # Fallback mais amplo: tudo a partir do segundo CERTIFICO
        m_segundo = re.search(
            r'(CERTIFICO,\s+ainda.+)',
            texto_flat, re.IGNORECASE | re.DOTALL
        )
    if m_segundo:
        trecho_resp = m_segundo.group(1)

    # Padrão: NOME (CPF/CNPJ: número)
    padrao = re.compile(
        r'(?:Sr\.?\s*)?'
        r'([A-ZÁÉÍÓÚÀÂÊÔÇÃẼÕÜ][A-ZÁÉÍÓÚÀÂÊÔÇÃẼÕÜA-Z\s\./]{2,}?)'
        r'\s*\((CPF|CNPJ)\s*:?\s*([\d\.\/\-]+)\)',
        re.IGNORECASE
    )

    vistos = set()
    for m in padrao.finditer(trecho_resp):
        nome    = m.group(1).strip()
        tipo    = m.group(2).upper()
        num_doc = m.group(3).strip()

        if num_doc in vistos:
            continue
        vistos.add(num_doc)

        # Remove prefixos indesejados: "e ", "Sr. ", "a " etc.
        nome = re.sub(r'^(?:e\s+|Sr\.?\s*|a\s+)', '', nome, flags=re.IGNORECASE).strip()

        # Busca o acórdão deste CPF/CNPJ no primeiro parágrafo
        # (onde estão os acórdãos individuais 185/2023, 185-B/2023, etc.)
        acordao = _buscar_acordao_por_doc(texto_flat, num_doc)

        dados["responsaveis"].append({
            "nome":       nome,
            "tipo_doc":   tipo,
            "numero_doc": num_doc,
            "acordao":    acordao,
            "excluido":   False   # por definição: só chegam aqui os remanescentes
        })

    # -------------------------------------------------------
    # 6) Fallback A: segundo parágrafo existe mas não tinha CPF/CNPJ
    #    (caso "os responsáveis acima citados" — sem repetir os nomes)
    #    → extrai do PRIMEIRO parágrafo, filtrando excluídos
    # Fallback B: sem marcadores a)/b)/c) — padrão "NOME (CPF/CNPJ: X)"
    #    direto no texto, filtrando excluídos
    # -------------------------------------------------------
    if not dados["responsaveis"]:
        # Padrão geral sem marcador de lista
        padrao_geral = re.compile(
            r'(?:Sr\.?\s*)?'
            r'([A-ZÁÉÍÓÚÀÂÊÔÇÃẼÕÜ][A-ZÁÉÍÓÚÀÂÊÔÇÃẼÕÜA-Z\s\./]{2,}?)'
            r'\s*\((CPF|CNPJ)\s*:?\s*([\d\.\/\-]+)\)',
            re.IGNORECASE
        )
        for m in padrao_geral.finditer(texto_flat):
            num_doc = m.group(3).strip()
            if num_doc in vistos:
                continue

            # Verifica exclusão nos ~700 chars seguintes
            trecho_pos = texto_flat[m.end(): m.end() + 700]
            excluido   = bool(re.search(
                r'excluiu[^.]{0,100}rol|excluído\s+do\s+rol',
                trecho_pos, re.IGNORECASE
            ))
            if excluido:
                continue

            vistos.add(num_doc)
            nome = re.sub(r'^(?:e\s+|Sr\.?\s*|a\s+)', '',
                          m.group(1).strip(), flags=re.IGNORECASE).strip()
            acordao = _buscar_acordao_por_doc(texto_flat, num_doc)
            dados["responsaveis"].append({
                "nome":       nome,
                "tipo_doc":   m.group(2).upper(),
                "numero_doc": num_doc,
                "acordao":    acordao,
                "excluido":   False
            })

    # Fallback C: se ainda vazio, tenta marcadores a)/b)/c) com filtro de exclusão
    if not dados["responsaveis"]:
        padrao_lista = re.compile(
            r'[a-z]\)\s*(?:Sr\.?\s*)?'
            r'([A-ZÁÉÍÓÚÀÂÊÔÇÃẼÕÜ][A-ZÁÉÍÓÚÀÂÊÔÇÃẼÕÜA-Z\s\./]+?)'
            r'\s*\((CPF|CNPJ)\s*:?\s*([\d\.\/\-]+)\)',
            re.IGNORECASE
        )
        for m in padrao_lista.finditer(texto_flat):
            num_doc = m.group(3).strip()
            if num_doc in vistos:
                continue
            trecho_pos = texto_flat[m.end(): m.end() + 700]
            excluido   = bool(re.search(
                r'excluiu[^.]{0,100}rol|excluído\s+do\s+rol',
                trecho_pos, re.IGNORECASE
            ))
            if excluido:
                continue
            vistos.add(num_doc)
            acordao = _buscar_acordao_por_doc(texto_flat, num_doc)
            dados["responsaveis"].append({
                "nome":       m.group(1).strip(),
                "tipo_doc":   m.group(2).upper(),
                "numero_doc": num_doc,
                "acordao":    acordao,
                "excluido":   False
            })

    return dados


# def _buscar_acordao_por_doc(texto_flat: str, num_doc: str) -> str | None:
#     """
#     Localiza o acórdão associado a um CPF/CNPJ no texto completo.
#     Busca o primeiro 'conforme Acórdão nº XXX' nos ~400 chars após o documento.
#     """
#     m_doc = re.search(re.escape(num_doc), texto_flat)
#     if not m_doc:
#         return None
#     trecho = texto_flat[m_doc.end(): m_doc.end() + 400]
#     # --- Acórdão ---
#     m_ac = re.search(
#         r"AC[ÓO]RD[ÃA]O\s*(?:N[º°]\s*)?[\s:]*\n*\s*(\d{1,5}/\d{4})",
#         texto,
#         flags=re.IGNORECASE
#     )

#     if not m_ac:
#         # fallback: captura qualquer número no padrão 000/0000
#         m_ac = re.search(r"\b\d{1,5}/\d{4}\b", texto)

#     acordao = m_ac.group(1) if m_ac and m_ac.lastindex else (
#         m_ac.group(0) if m_ac else ""
#     )
   
#     return _normalizar_acordao(acordao if m_ac else None

def _buscar_acordao_por_doc(texto_flat: str, num_doc: str) -> str | None:
    """
    Localiza o acórdão associado a um CPF/CNPJ no texto completo.
    Busca o primeiro 'Acórdão XXX/AAAA' nos ~400 chars após o documento.
    """
    m_doc = re.search(re.escape(num_doc), texto_flat)
    if not m_doc:
        return None

    trecho = texto_flat[m_doc.end(): m_doc.end() + 400]

    # procura após a palavra acórdão (mais robusto)
    m_ac = re.search(
        r"AC[ÓO]RD[ÃA]O.*?(\d{1,5}/\d{4})",
        trecho,
        flags=re.IGNORECASE | re.DOTALL
    )

    if not m_ac:
        # fallback: qualquer número no padrão 000/0000 dentro do trecho
        m_ac = re.search(r"\b\d{1,5}/\d{4}\b", trecho)

    if not m_ac:
        return None

    return _normalizar_acordao(m_ac.group(1) if m_ac.lastindex else m_ac.group(0))
# ============================================================
# BANCO DE DADOS SQLite
# ============================================================

def inicializar_banco(caminho_db: str = "certidoes_tce.db") -> sqlite3.Connection:
    """Cria / abre o banco SQLite e garante que as tabelas existam."""
    conn = sqlite3.connect(caminho_db)
    cursor = conn.cursor()

    cursor.executescript("""
    CREATE TABLE IF NOT EXISTS certidoes (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        numero_processo TEXT NOT NULL,
        valor_atualizado TEXT,
        data_atualizacao TEXT,
        acordao_origem   TEXT,
        data_insercao    TEXT DEFAULT CURRENT_TIMESTAMP,
        caminho_pdf_certidao TEXT,
        caminho_pdf_planilha  TEXT
    );

    CREATE TABLE IF NOT EXISTS responsaveis (
        id               INTEGER PRIMARY KEY AUTOINCREMENT,
        certidao_id      INTEGER NOT NULL REFERENCES certidoes(id),
        nome             TEXT NOT NULL,
        tipo_doc         TEXT,   -- CPF ou CNPJ
        numero_doc       TEXT,
        acordao          TEXT,
        excluido         INTEGER DEFAULT 0,
        endereco         TEXT
    );
    """)
    conn.commit()
    return conn


def salvar_certidao(conn: sqlite3.Connection, dados: dict,
                     caminho_certidao: str = None,
                     caminho_planilha: str = None) -> int:
    """Salva os dados extraídos no banco. Retorna o ID da certidão inserida."""
    cursor = conn.cursor()

    # Verifica se o processo já existe
    cursor.execute(
        "SELECT id FROM certidoes WHERE numero_processo = ?",
        (dados["numero_processo"],)
    )
    row = cursor.fetchone()
    if row:
        certidao_id = row[0]
        cursor.execute("""
            UPDATE certidoes SET
                valor_atualizado = ?,
                data_atualizacao = ?,
                acordao_origem   = ?,
                caminho_pdf_certidao = ?,
                caminho_pdf_planilha  = ?
            WHERE id = ?
        """, (dados["valor_atualizado"], dados["data_atualizacao"],
              dados["acordao_origem"], caminho_certidao, caminho_planilha,
              certidao_id))
        # Remove responsáveis antigos para reinserir
        cursor.execute("DELETE FROM responsaveis WHERE certidao_id = ?",
                       (certidao_id,))
    else:
        cursor.execute("""
            INSERT INTO certidoes
                (numero_processo, valor_atualizado, data_atualizacao,
                 acordao_origem, caminho_pdf_certidao, caminho_pdf_planilha)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (dados["numero_processo"], dados["valor_atualizado"],
              dados["data_atualizacao"], dados["acordao_origem"],
              caminho_certidao, caminho_planilha))
        certidao_id = cursor.lastrowid

    for resp in dados["responsaveis"]:
        cursor.execute("""
            INSERT INTO responsaveis
                (certidao_id, nome, tipo_doc, numero_doc, acordao, excluido)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (certidao_id, resp["nome"], resp["tipo_doc"],
              resp["numero_doc"], resp["acordao"],
              1 if resp["excluido"] else 0))

    conn.commit()
    print(f"✅ Certidão {dados['numero_processo']} salva no banco (ID={certidao_id})")
    return certidao_id


def listar_certidoes(conn: sqlite3.Connection):
    """Retorna lista de todas as certidões com seus responsáveis."""
    cursor = conn.cursor()
    cursor.execute("""
        SELECT c.id, c.numero_processo, c.valor_atualizado, c.data_atualizacao,
               c.acordao_origem
        FROM certidoes c
        ORDER BY c.id DESC
    """)
    certidoes = cursor.fetchall()

    resultado = []
    for c in certidoes:
        cursor.execute("""
            SELECT nome, tipo_doc, numero_doc, acordao, excluido, endereco
            FROM responsaveis WHERE certidao_id = ?
        """, (c[0],))
        resps = cursor.fetchall()
        resultado.append({
            "id": c[0],
            "numero_processo": c[1],
            "valor_atualizado": c[2],
            "data_atualizacao": c[3],
            "acordao_origem": c[4],
            "responsaveis": [
                {"nome": r[0], "tipo_doc": r[1], "numero_doc": r[2],
                 "acordao": r[3], "excluido": bool(r[4]), "endereco": r[5]}
                for r in resps
            ]
        })
    return resultado


# ============================================================
# GERAÇÃO DA PETIÇÃO (DOCX via Node/docx-js)
# ============================================================

def gerar_valor_extenso(valor_str: str) -> str:
    """
    Converte 'R$ 167.406,32' → 'cento e sessenta e sete mil e quatrocentos
    e seis reais e trinta e dois centavos'.
    100% Python puro — sem Node.js, sem dependências externas.
    """
    # Normaliza: remove "R$", tira pontos de milhar, troca vírgula por ponto
    limpo = re.sub(r'[R$\s]', '', valor_str).replace('.', '').replace(',', '.')
    try:
        n = float(limpo)
    except ValueError:
        return valor_str

    inteiros  = int(n)
    centavos  = round((n - inteiros) * 100)

    UNIDADES = ["", "um", "dois", "três", "quatro", "cinco", "seis",
                "sete", "oito", "nove", "dez", "onze", "doze", "treze",
                "quatorze", "quinze", "dezesseis", "dezessete", "dezoito",
                "dezenove"]
    DEZENAS  = ["", "", "vinte", "trinta", "quarenta", "cinquenta",
                "sessenta", "setenta", "oitenta", "noventa"]
    CENTENAS = ["", "cem", "duzentos", "trezentos", "quatrocentos",
                "quinhentos", "seiscentos", "setecentos", "oitocentos",
                "novecentos"]

    def _extenso(n: int) -> str:
        if n == 0:
            return "zero"
        if n == 100:
            return "cem"
        if n < 20:
            return UNIDADES[n]
        if n < 100:
            d, u = divmod(n, 10)
            return DEZENAS[d] + (" e " + UNIDADES[u] if u else "")
        if n < 1000:
            c, r = divmod(n, 100)
            base = "cento" if c == 1 and r > 0 else CENTENAS[c]
            return base + (" e " + _extenso(r) if r else "")
        if n < 1_000_000:
            m, r = divmod(n, 1000)
            prefixo = "mil" if m == 1 else _extenso(m) + " mil"
            return prefixo + (" e " + _extenso(r) if r else "")
        if n < 1_000_000_000:
            mi, r = divmod(n, 1_000_000)
            sufixo = "milhão" if mi == 1 else "milhões"
            parte_mi = _extenso(mi) + " " + sufixo
            return parte_mi + (" e " + _extenso(r) if r else "")
        return str(n)

    partes = []
    if inteiros > 0:
        partes.append(_extenso(inteiros) + (" real" if inteiros == 1 else " reais"))
    if centavos > 0:
        partes.append(_extenso(centavos) + (" centavo" if centavos == 1 else " centavos"))

    return " e ".join(partes) if partes else "zero reais"


def gerar_peticao_docx(dados: dict, caminho_saida: str,
                        data_peticao: str = None,
                        procurador_nome: str = "Cid Carlos Gonçalves Coelho",
                        procurador_oab: str = "OAB/PI 2844") -> str:
    """
    Gera a petição inicial de execução extrajudicial em DOCX.
    100% Python puro via python-docx — sem Node.js.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    # ── helpers ──────────────────────────────────────────────────────────────
    def _set_cell_color(cell, hex_color: str):
        """Preenche fundo de célula com cor hex (ex: '1A3A5C')."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def _set_cell_borders(cell, color="2C5F8A", sz="6"):
        """Aplica bordas simples em todos os lados de uma célula."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    sz)
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), color)
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def _cell_margins(cell, top=80, bottom=80, left=120, right=120):
        """Define margens internas de célula em twips."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top', top), ('bottom', bottom),
                          ('left', left), ('right', right)]:
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'),    str(val))
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)

    def _add_run(para, text, bold=False, italic=False,
                 underline=False, font_size=12, color=None):
        run = para.add_run(text)
        run.bold      = bold
        run.italic    = italic
        run.underline = underline
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        return run

    def _para(doc, text="", bold=False, italic=False, underline=False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, font_size=12,
              space_before=0, space_after=6, line_spacing=None):
        p = doc.add_paragraph()
        p.alignment = align
        fmt = p.paragraph_format
        fmt.space_before = Pt(space_before)
        fmt.space_after  = Pt(space_after)
        if line_spacing:
            from docx.shared import Pt as _Pt
            fmt.line_spacing = _Pt(line_spacing)
        if text:
            _add_run(p, text, bold=bold, italic=italic,
                     underline=underline, font_size=font_size)
        return p

    # def _mixed_para(doc, runs, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    #                 space_before=0, space_after=6):
    #     """runs: lista de dicts com keys: text, bold, italic, underline, size."""
    #     p = doc.add_paragraph()
    #     p.alignment = align
    #     p.paragraph_format.space_before = Pt(space_before)
    #     p.paragraph_format.space_after  = Pt(space_after)
    #     for r in runs:
    #         _add_run(p,
    #                  text=r.get("text", ""),
    #                  bold=r.get("bold", False),
    #                  italic=r.get("italic", False),
    #                  underline=r.get("underline", False),
    #                  font_size=r.get("size", 12))
    #     return p
    
    def _mixed_para(doc, runs, space_before=0, space_after=0):
        paragraph = doc.add_paragraph()

        # ── Formatação padrão da petição ─────────────────────────────
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(2)
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.line_spacing = 1.5

        # ── Se vier string simples ────────────────────────────────────
        if isinstance(runs, str):
            paragraph.add_run(runs)
            return paragraph

        # ── Se vier lista estruturada ─────────────────────────────────
        for item in runs:
            if isinstance(item, str):
                run = paragraph.add_run(item)
            else:
                run = paragraph.add_run(item.get("text", ""))
                run.bold = item.get("bold", False)
                run.italic = item.get("italic", False)
                run.underline = item.get("underline", False)

        return paragraph

    # ── data de petição ───────────────────────────────────────────────────────
    MESES_PT = {
        "January": "janeiro", "February": "fevereiro", "March": "março",
        "April": "abril",     "May": "maio",           "June": "junho",
        "July": "julho",      "August": "agosto",      "September": "setembro",
        "October": "outubro", "November": "novembro",  "December": "dezembro",
    }
    if data_peticao is None:
        data_peticao = datetime.now().strftime("%d de %B de %Y")
        for en, pt in MESES_PT.items():
            data_peticao = data_peticao.replace(en, pt)

    # ── dados ─────────────────────────────────────────────────────────────────
    # Tabela mostra TODOS os responsáveis; cabeçalho lista apenas os ativos
    todos_responsaveis  = dados["responsaveis"]
    responsaveis_ativos = [r for r in todos_responsaveis if not r["excluido"]]
    nomes_executados    = ", ".join(r["nome"] for r in responsaveis_ativos)
    valor_atualizado    = dados["valor_atualizado"] or "R$ 0,00"
    data_atualizacao    = dados["data_atualizacao"] or ""
    numero_processo     = dados["numero_processo"]  or ""
    valor_extenso       = gerar_valor_extenso(valor_atualizado)

    # ── documento ─────────────────────────────────────────────────────────────
    doc = Document()

    # Margens da página (A4)
    for section in doc.sections:
        section.page_width  = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Fonte padrão do documento
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    C = WD_ALIGN_PARAGRAPH.CENTER
    J = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Cabeçalho (JUSTIFICADO conforme solicitado) ───────────────────────────
    _para(doc,
          "EXMO. SR. DR. JUIZ DE DIREITO DA VARA DA FAZENDA PÚBLICA DA COMARCA DE TERESINA - PI",
          bold=True, align=J, space_after=30)
    # _para(doc, "Execução de Título Extrajudicial", align=J, space_after=2)
    _para(doc, "Exequente: Estado do Piauí",       align=J, space_after=2)
    _para(doc, f"Executado: {nomes_executados}.",  align=J, space_after=90)

    # ── Qualificação ──────────────────────────────────────────────────────────
    _mixed_para(doc, [
        {"text": "O ESTADO DO PIAUÍ,", "bold": True},
        {"text": (" pessoa jurídica de Direito Público interno, representado em juízo por seus "
                  "procuradores (conforme os artigos 132, da Constituição da República; 150 da "
                  "Constituição do Estado do Piauí; 75, II, do Novo Código de Processo Civil/2015 "
                  "e 2º da Lei Complementar Estadual n. 56/2005), com endereço para comunicações "
                  "processuais na Avenida Senador Arêa Leão, n. 1650, Jóquei, Teresina (PI), vem "
                  "respeitosamente à presença de Vossa Excelência, ajuizar nos termos dos arts. "
                  "71, § 3º, da Constituição Federal, art. 86, § 2º, da Constituição Estadual, "
                  "art. 135, 140, da Lei n.º 5.888/2009 (Lei Orgânica do TCE), bem como nos "
                  "artigos 783 e seguintes do Novo Código de Processo Civil "
                  "(Lei n.º 13.205/2015), a presente")}
    ], space_after=8)

    _para(doc, "EXECUÇÃO POR TÍTULO EXTRAJUDICIAL",
          bold=True, align=C, space_after=8)
    _para(doc, "em desfavor de", align=C, space_after=10)

    # ── Tabela de executados (TODOS, inclusive excluídos com nota) ────────────
    HDR_LABELS = ["EXECUTADO", "CPF/CNPJ", "ENDEREÇO", "Nº PROCESSO", "ACÓRDÃO TCE"]
    # Larguras das colunas em cm (total ≈ 16 cm de área útil)
    COL_CM = [4.5, 3.2, 3.0, 2.8, 2.5]

    n_cols = len(HDR_LABELS)
    n_rows = 1 + len(todos_responsaveis) + 1   # header + todos + linha vazia

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Larguras das colunas
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(COL_CM[i])

    # Linha de cabeçalho
    for i, label in enumerate(HDR_LABELS):
        cell = tbl.rows[0].cells[i]
        _set_cell_color(cell, "1A3A5C")
        _set_cell_borders(cell, "1A3A5C", "8")
        _cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = C
        r = p.add_run(label)
        r.bold       = True
        r.font.name  = "Times New Roman"
        r.font.size  = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Linhas de dados — todos os responsáveis (sem distinção visual para excluídos)
    for idx, resp in enumerate(todos_responsaveis):
        row = tbl.rows[1 + idx]
        valores = [
            resp["nome"],
            f"{resp['tipo_doc']}: {resp['numero_doc']}",
            resp.get("endereco") or "",
            numero_processo,
            resp["acordao"] or dados["acordao_origem"] or ""
        ]
        for i, val in enumerate(valores):
            cell = row.cells[i]
            _set_cell_borders(cell)
            _cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = C
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)

    # Linha vazia final
    row_vazia = tbl.rows[1 + len(todos_responsaveis)]
    for i in range(n_cols):
        cell = row_vazia.cells[i]
        _set_cell_borders(cell)
        _cell_margins(cell)

    doc.add_paragraph()   # espaço após tabela

    # ── Objeto da execução ────────────────────────────────────────────────────
    _mixed_para(doc, [
        {"text": ("A presente execução por título extrajudicial tem por objeto a Certidão "
                  "de Débito referente à Imputação de Débito aos executados, constituída na "
                  "forma dos arts. 71, § 3º da Constituição Federal, art. 86, § 2º da "
                  "Constituição Estadual, art. 135, 140 da Lei n.º 5.888/2009, cujo valor foi "
                  f"atualizado monetariamente até {data_atualizacao}, referente à imputação "
                  "de débito, o qual devem restituir aos cofres do Estado do Piauí o valor de ")},
        {"text": valor_atualizado, "bold": True},
        {"text": (", que deverá ser atualizado pela SELIC até a data do efetivo pagamento, "
                  "conforme o disposto no art. 140, caput, da Lei Estadual nº 5.888/2009 "
                  "(Lei Orgânica do TCE/PI) bem como na Lei Estadual nº 6.782/2016, que "
                  "instituiu a SELIC como índice de atualização no âmbito do Estado do Piauí, "
                  "ou por outro indexador que venha a substituí-la.")}
    ], space_before=6, space_after=6)

    # ── Base legal ────────────────────────────────────────────────────────────
    _mixed_para(doc, [
        {"text": ("Vale destacar que a presente execução por título extrajudicial deve "
                  "seguir os termos ")},
        {"text": ("dos arts. 778, caput, 779, 781, 783, 784, 786, 798 e 829 do Novo "
                  "Código Processo Civil/2015"),
         "bold": True, "underline": True},
        {"text": ", bem como as disposições aplicáveis à espécie."}
    ], space_after=6)

    # ── Pedidos ───────────────────────────────────────────────────────────────
    _mixed_para(doc, [
        {"text": ("À vista do exposto, considerando-se a liquidez, certeza e exigibilidade "
                  "do título anexo, requer a CITAÇÃO dos executados, nos endereços indicados "
                  "na tabela acima, por carta e com aviso de recebimento, ou por mandado a "
                  "ser cumprido pelo oficial de justiça, para no prazo de 03 (três) dias, "
                  "efetuarem ")},
        {"text": "SOLIDARIAMENTE", "bold": True},
        {"text": (" o pagamento do débito (NCPC/2015, art. 829), acrescido de todos os ônus "
                  "de sucumbência, atualização monetária, custas, despesas processuais, com "
                  "condenação do devedor no pagamento de honorários advocatícios no percentual "
                  "de 20 % sobre o valor da causa devidamente atualizado.")}
    ], space_after=6)

    _mixed_para(doc,
          ("Na hipótese de não pagamento integral do débito devidamente atualizado, "
           "proceda-se à penhora de bens do devedor passíveis de constrição, priorizando-se "
           "o dinheiro (art. 835, I e 854 do NCPC/2015), o que poderá ser feito mediante "
           "acesso ao sistema SISBAJUD, que está à disposição do MM. Juízo;"),
          space_after=6)

    _mixed_para(doc,
          ("Destacamos que o oficial de justiça, não encontrando o devedor, arrestar-lhe-á "
           "tantos bens quantos bastem para garantir a execução (CPC, art. 830), dando-se "
           "ciência ao exequente para as providências legais."),
          space_after=6)

    _mixed_para(doc,
          ("Requer, feita a penhora sobre bens imóveis, a devida averbação da penhora no "
           "respectivo Registro de Imóveis, observando o disposto no art. 828, do NCPC/2015."),
          space_after=6)

    _mixed_para(doc,
          ("Ademais, requer a expropriação dos bens penhorados e o final pagamento do débito, "
           "devidamente acrescido de todos os ônus de sucumbência, juros de mora, custas e "
           "despesas processuais, com condenação do devedor no pagamento de honorários "
           "advocatícios no percentual de 20% sobre o valor da causa devidamente atualizado."),
          space_after=10)

    # ── Valor da causa ────────────────────────────────────────────────────────
    _mixed_para(doc, [
        {"text": "Dá à causa o valor de "},
        {"text": f"{valor_atualizado} ({valor_extenso}).",
         "bold": True, "underline": True},
    ], space_after=18)

    # ── Fecho ─────────────────────────────────────────────────────────────────
    _para(doc, f"Teresina, {data_peticao}.", space_after=18)
    _para(doc, procurador_nome,  align=C, space_after=2)
    _para(doc, "Procurador do Estado", align=C, space_after=2)
    _para(doc, procurador_oab,   align=C)

    doc.save(caminho_saida)
    return caminho_saida


# ============================================================
# GERAÇÃO DO PDF FINAL (petição + certidão + planilha)
# ============================================================

def gerar_peticao_pdf(dados: dict, caminho_pdf: str,
                      data_peticao: str = None,
                      procurador_nome: str = "Cid Carlos Gonçalves Coelho",
                      procurador_oab: str = "OAB/PI 2844") -> str:
    """
    Gera a petição diretamente em PDF usando reportlab.
    100% Python puro — sem LibreOffice, sem Word, sem Node.js.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
    )

    MESES_PT = {
        "January": "janeiro", "February": "fevereiro", "March": "março",
        "April": "abril",     "May": "maio",           "June": "junho",
        "July": "julho",      "August": "agosto",      "September": "setembro",
        "October": "outubro", "November": "novembro",  "December": "dezembro",
    }
    if data_peticao is None:
        data_peticao = datetime.now().strftime("%d de %B de %Y")
        for en, pt in MESES_PT.items():
            data_peticao = data_peticao.replace(en, pt)

    responsaveis_ativos = [r for r in dados["responsaveis"] if not r["excluido"]]
    nomes_executados    = ", ".join(r["nome"] for r in responsaveis_ativos)
    valor_atualizado    = dados["valor_atualizado"] or "R$ 0,00"
    data_atualizacao    = dados["data_atualizacao"] or ""
    numero_processo     = dados["numero_processo"]  or ""
    valor_extenso       = gerar_valor_extenso(valor_atualizado)

    # ── Estilos ────────────────────────────────────────────────────────────────
    FONTE   = "Times-Roman"
    FONTE_B = "Times-Bold"
    FONTE_BI = "Times-BoldItalic"
    TAM  = 12
    TAM_S = 10   # tabela
    
    s_normal_sem_recuo = ParagraphStyle("normal",
        fontName=FONTE, fontSize=TAM,
        alignment=TA_JUSTIFY, leading=TAM * 1.5,
        spaceAfter=4)

    s_normal = ParagraphStyle("normal",
        fontName=FONTE, fontSize=TAM,
        alignment=TA_JUSTIFY, leading=TAM * 1.5,
        spaceAfter=4,
        firstLineIndent=2 * cm)

    s_center = ParagraphStyle("center",
        fontName=FONTE, fontSize=TAM,
        alignment=TA_CENTER, leading=TAM * 1.5,
        spaceAfter=4)

    s_bold_center = ParagraphStyle("bold_center",
        fontName=FONTE_B, fontSize=TAM,
        alignment=TA_CENTER, leading=TAM * 1.5,
        spaceAfter=4)

    s_bold_just = ParagraphStyle("bold_just",
        fontName=FONTE_B, fontSize=TAM,
        alignment=TA_JUSTIFY, leading=TAM * 1.5,
        spaceAfter=4)

    def p(text, style=s_normal):
        return Paragraph(text, style)

    def sp(h=6):
        return Spacer(1, h)

    # ── Conteúdo ───────────────────────────────────────────────────────────────
    story = []

    # Cabeçalho (justificado)
    story.append(p("<b>EXMO. SR. DR. JUIZ DE DIREITO DA VARA DA FAZENDA PÚBLICA "
                   "DA COMARCA DE TERESINA - PI</b>", s_bold_just))
    story.append(sp(32))
    # story.append(p("Execução de Título Extrajudicial", s_normal_sem_recuo))
    story.append(p("Exequente: Estado do Piauí",       s_normal_sem_recuo))
    story.append(p(f"Executado: {nomes_executados}.",  s_normal_sem_recuo))
    story.append(sp(64))

    # Qualificação
    story.append(p(
        f"<b>O ESTADO DO PIAUÍ,</b> pessoa jurídica de Direito Público interno, "
        "representado em juízo por seus procuradores (conforme os artigos 132, da "
        "Constituição da República; 150 da Constituição do Estado do Piauí; 75, II, "
        "do Novo Código de Processo Civil/2015 e 2º da Lei Complementar Estadual "
        "n. 56/2005), com endereço para comunicações processuais na Avenida Senador "
        "Arêa Leão, n. 1650, Jóquei, Teresina (PI), vem respeitosamente à presença "
        "de Vossa Excelência, ajuizar nos termos dos arts. 71, § 3º, da Constituição "
        "Federal, art. 86, § 2º, da Constituição Estadual, art. 135, 140, da Lei "
        "n.º 5.888/2009 (Lei Orgânica do TCE), bem como nos artigos 783 e seguintes "
        "do Novo Código de Processo Civil (Lei n.º 13.205/2015), a presente"))
    story.append(sp(8))

    story.append(p("<b>EXECUÇÃO POR TÍTULO EXTRAJUDICIAL</b>", s_bold_center))
    story.append(p("em desfavor de", s_center))
    story.append(sp(8))

    # ── Tabela ─────────────────────────────────────────────────────────────────
    HDR = ["EXECUTADO", "CPF/CNPJ", "ENDEREÇO", "Nº PROCESSO", "ACÓRDÃO TCE"]
    # Larguras proporcionais à área útil (A4 com margens 3/2 cm → ~16 cm)
    area = 16 * cm
    col_w = [area * p for p in [0.28, 0.20, 0.19, 0.18, 0.15]]

    # Linha de cabeçalho
    hdr_row = [Paragraph(f"<b><font color='white'>{h}</font></b>",
                         ParagraphStyle("th", fontName=FONTE_B, fontSize=TAM_S,
                                        alignment=TA_CENTER, leading=TAM_S*1.3))
               for h in HDR]

    # Linhas de dados
    data_rows = []
    for resp in responsaveis_ativos:
        data_rows.append([
            Paragraph(resp["nome"],
                      ParagraphStyle("td", fontName=FONTE, fontSize=TAM_S,
                                     alignment=TA_CENTER, leading=TAM_S*1.3)),
            Paragraph(f"{resp['tipo_doc']}: {resp['numero_doc']}",
                      ParagraphStyle("td", fontName=FONTE, fontSize=TAM_S,
                                     alignment=TA_CENTER, leading=TAM_S*1.3)),
            Paragraph(resp.get("endereco") or "",
                      ParagraphStyle("td", fontName=FONTE, fontSize=TAM_S,
                                     alignment=TA_CENTER, leading=TAM_S*1.3)),
            Paragraph(numero_processo,
                      ParagraphStyle("td", fontName=FONTE, fontSize=TAM_S,
                                     alignment=TA_CENTER, leading=TAM_S*1.3)),
            Paragraph(resp["acordao"] or dados["acordao_origem"] or "",
                      ParagraphStyle("td", fontName=FONTE, fontSize=TAM_S,
                                     alignment=TA_CENTER, leading=TAM_S*1.3)),
        ])
    # Linha vazia
    data_rows.append(["", "", "", "", ""])

    tbl_data = [hdr_row] + data_rows
    tbl = Table(tbl_data, colWidths=col_w)
    COR_HDR = colors.HexColor("#1A3A5C")
    COR_BRD = colors.HexColor("#2C5F8A")
    tbl.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0), (-1, 0),  COR_HDR),
        ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
        ("GRID",        (0, 0), (-1, -1), 0.5, COR_BRD),
        ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",  (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",(0,0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING",(0, 0), (-1, -1), 4),
    ]))
    story.append(KeepTogether(tbl))
    story.append(sp(10))

    # ── Corpo da petição ───────────────────────────────────────────────────────
    story.append(p(
        f"A presente execução por título extrajudicial tem por objeto a Certidão de "
        "Débito referente à Imputação de Débito aos executados, constituída na forma "
        "dos arts. 71, § 3º da Constituição Federal, art. 86, § 2º da Constituição "
        "Estadual, art. 135, 140 da Lei n.º 5.888/2009, cujo valor foi atualizado "
        f"monetariamente até {data_atualizacao}, referente à imputação de débito, o "
        "qual devem restituir aos cofres do Estado do Piauí o valor de "
        f"<b>{valor_atualizado}</b>, que deverá ser atualizado pela SELIC até a data "
        "do efetivo pagamento, conforme o disposto no art. 140, caput, da Lei Estadual "
        "nº 5.888/2009 (Lei Orgânica do TCE/PI) bem como na Lei Estadual nº 6.782/2016, "
        "que instituiu a SELIC como índice de atualização no âmbito do Estado do Piauí, "
        "ou por outro indexador que venha a substituí-la."))
    story.append(sp(4))

    story.append(p(
        "Vale destacar que a presente execução por título extrajudicial deve seguir os "
        "termos <b><u>dos arts. 778, caput, 779, 781, 783, 784, 786, 798 e 829 do Novo "
        "Código Processo Civil/2015</u></b>, bem como as disposições aplicáveis à espécie."))
    story.append(sp(4))

    story.append(p(
        "À vista do exposto, considerando-se a liquidez, certeza e exigibilidade do "
        "título anexo, requer a CITAÇÃO dos executados, nos endereços indicados na "
        "tabela acima, por carta e com aviso de recebimento, ou por mandado a ser "
        "cumprido pelo oficial de justiça, para no prazo de 03 (três) dias, efetuarem "
        "<b>SOLIDARIAMENTE</b> o pagamento do débito (NCPC/2015, art. 829), acrescido "
        "de todos os ônus de sucumbência, atualização monetária, custas, despesas "
        "processuais, com condenação do devedor no pagamento de honorários advocatícios "
        "no percentual de 20 % sobre o valor da causa devidamente atualizado."))
    story.append(sp(4))

    story.append(p(
        "Na hipótese de não pagamento integral do débito devidamente atualizado, "
        "proceda-se à penhora de bens do devedor passíveis de constrição, "
        "priorizando-se o dinheiro (art. 835, I e 854 do NCPC/2015), o que poderá "
        "ser feito mediante acesso ao sistema SISBAJUD, que está à disposição do MM. Juízo;"))
    story.append(sp(4))

    story.append(p(
        "Destacamos que o oficial de justiça, não encontrando o devedor, "
        "arrestar-lhe-á tantos bens quantos bastem para garantir a execução "
        "(CPC, art. 830), dando-se ciência ao exequente para as providências legais."))
    story.append(sp(4))

    story.append(p(
        "Requer, feita a penhora sobre bens imóveis, a devida averbação da penhora "
        "no respectivo Registro de Imóveis, observando o disposto no art. 828, do NCPC/2015."))
    story.append(sp(4))

    story.append(p(
        "Ademais, requer a expropriação dos bens penhorados e o final pagamento do "
        "débito, devidamente acrescido de todos os ônus de sucumbência, juros de mora, "
        "custas e despesas processuais, com condenação do devedor no pagamento de "
        "honorários advocatícios no percentual de 20% sobre o valor da causa "
        "devidamente atualizado."))
    story.append(sp(10))

    # Valor da causa
    story.append(p(
        f"Dá à causa o valor de <b><u>{valor_atualizado} ({valor_extenso}).</u></b>"))
    story.append(sp(16))

    # Fecho
    story.append(p(f"Teresina, {data_peticao}."))
    story.append(sp(16))
    story.append(p(procurador_nome,      s_center))
    story.append(p("Procurador do Estado", s_center))
    story.append(p(procurador_oab,        s_center))

    # ── Monta o documento ──────────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        caminho_pdf,
        pagesize=A4,
        leftMargin=3*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )
    doc.build(story)
    return caminho_pdf


def converter_docx_para_pdf(caminho_docx: str, caminho_pdf: str) -> str:
    """
    Tentativa de conversão DOCX→PDF via ferramentas externas.
    Tenta LibreOffice, Word COM e docx2pdf.
    Se nenhum funcionar, lança RuntimeError — use gerar_peticao_pdf() como alternativa.
    """
    import shutil

    pasta_saida = str(Path(caminho_pdf).parent)

    for exe in ("libreoffice", "soffice",
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"):
        try:
            result = subprocess.run(
                [exe, "--headless", "--convert-to", "pdf",
                 "--outdir", pasta_saida, caminho_docx],
                capture_output=True, text=True, timeout=60)
            pdf_gerado = str(Path(pasta_saida) / (Path(caminho_docx).stem + ".pdf"))
            if os.path.exists(pdf_gerado):
                if pdf_gerado != caminho_pdf:
                    shutil.move(pdf_gerado, caminho_pdf)
                return caminho_pdf
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    if platform.system() == "Windows":
        try:
            import comtypes.client
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            doc_word = word.Documents.Open(str(Path(caminho_docx).resolve()))
            doc_word.SaveAs(str(Path(caminho_pdf).resolve()), FileFormat=17)
            doc_word.Close()
            word.Quit()
            if os.path.exists(caminho_pdf):
                return caminho_pdf
        except Exception:
            pass

    try:
        from docx2pdf import convert
        convert(caminho_docx, caminho_pdf)
        if os.path.exists(caminho_pdf):
            return caminho_pdf
    except Exception:
        pass

    raise RuntimeError("Nenhum conversor externo disponível.")


def montar_pdf_final(caminho_peticao_pdf: str,
                     caminho_certidao_pdf: str,
                     caminho_planilha_pdf: str,
                     caminho_saida: str) -> str:
    """Mescla petição + certidão + planilha num único PDF."""
    
    writer = PdfWriter()

    for pdf_path in [caminho_peticao_pdf, caminho_certidao_pdf, caminho_planilha_pdf]:
        if pdf_path and os.path.exists(pdf_path):
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                writer.add_page(page)

    with open(caminho_saida, "wb") as f:
        writer.write(f)

    # ✅ Exclui apenas a petição intermediária
    if caminho_peticao_pdf and os.path.exists(caminho_peticao_pdf):
        os.remove(caminho_peticao_pdf)

    return caminho_saida


# ============================================================
# INTERFACE DE LINHA DE COMANDO (CLI) SIMPLES
# ============================================================

def processar_certidao(
    caminho_certidao: str,
    caminho_planilha: str = None,
    pasta_saida: str = "saida_peticoes",
    caminho_db: str = "certidoes_tce.db",
    caminho_enderecos: str = None,       # ← PDF de endereços das partes (opcional)
) -> dict:
    """
    Fluxo completo:
      1. Extrai dados da certidão PDF
      2. Salva no banco SQLite
      3. (Opcional) Importa endereços do PDF de endereços → popula responsaveis
      4. Recarrega responsáveis do banco (com endereços já preenchidos)
      5. Gera petição DOCX
      6. Converte petição para PDF
      7. Mescla petição + certidão + planilha num único PDF final
    Retorna dict com caminhos gerados.
    """
    os.makedirs(pasta_saida, exist_ok=True)

    print(f"📄 Lendo certidão: {caminho_certidao}")
    texto_certidao = extrair_texto_pdf(caminho_certidao)

    print("🔍 Extraindo dados...")
    dados = extrair_dados_certidao(texto_certidao)

    print("\n📋 Dados extraídos:")
    print(f"  Processo:     {dados['numero_processo']}")
    print(f"  Valor:        {dados['valor_atualizado']}")
    print(f"  Atualizado:   {dados['data_atualizacao']}")
    print(f"  Acórdão orig: {dados['acordao_origem']}")
    for r in dados["responsaveis"]:
        status = " [EXCLUÍDO]" if r["excluido"] else ""
        print(f"  Responsável:  {r['nome']} | {r['tipo_doc']}: {r['numero_doc']} | "
              f"Acórdão: {r['acordao']}{status}")

    # ── Banco de dados ────────────────────────────────────────────────────────
    conn = inicializar_banco(caminho_db)
    certidao_id = salvar_certidao(conn, dados, caminho_certidao, caminho_planilha)

    # ── Importa endereços ANTES de gerar a petição ────────────────────────────
    if caminho_enderecos and os.path.isfile(caminho_enderecos):
        print(f"\n📍 Importando endereços: {os.path.basename(caminho_enderecos)}")
        try:
            from gestor_enderecos import processar_pdf_enderecos
            processar_pdf_enderecos(caminho_enderecos, conn,
                                    certidao_id=certidao_id, verbose=True)
        except ImportError:
            print("  ⚠️  gestor_enderecos.py não encontrado — endereços não importados.")
        except Exception as e_end:
            print(f"  ⚠️  Erro ao importar endereços: {e_end}")

    # ── Recarrega responsáveis do banco (agora com endereços preenchidos) ─────
    cursor = conn.cursor()
    cursor.execute(
        "SELECT nome, tipo_doc, numero_doc, acordao, excluido, endereco "
        "FROM responsaveis WHERE certidao_id = ? ORDER BY id",
        (certidao_id,)
    )
    rows = cursor.fetchall()
    if rows:
        dados["responsaveis"] = [
            {
                "nome":       r[0],
                "tipo_doc":   r[1],
                "numero_doc": r[2],
                "acordao":    r[3],
                "excluido":   bool(r[4]),
                "endereco":   r[5] or "",
            }
            for r in rows
        ]
        n_end = sum(1 for r in dados["responsaveis"] if r["endereco"])
        print(f"  ✅ {len(rows)} responsável(eis) recarregado(s) do banco "
              f"({n_end} com endereço)")

    conn.close()

    # ── Nomes de arquivo ──────────────────────────────────────────────────────
    proc_safe = (dados["numero_processo"] or "certidao").replace("/", "_").replace("\\", "_")
    docx_path      = os.path.join(pasta_saida, f"peticao_{proc_safe}.docx")
    pdf_pet_path   = os.path.join(pasta_saida, f"peticao_{proc_safe}.pdf")
    pdf_final_path = os.path.join(pasta_saida, f"EXECUCAO_{proc_safe}_COMPLETO.pdf")

    print(f"\n📝 Gerando petição DOCX: {docx_path}")
    gerar_peticao_docx(dados, docx_path)

    print(f"🖨️  Gerando petição PDF: {pdf_pet_path}")
    gerar_peticao_pdf(dados, pdf_pet_path)

    print(f"📦 Montando PDF final: {pdf_final_path}")
    montar_pdf_final(pdf_pet_path, caminho_certidao,
                     caminho_planilha, pdf_final_path)
    print(f"✅ PDF final gerado: {pdf_final_path}")

    return {
        "dados": dados,
        "certidao_id": certidao_id,
        "docx": docx_path,
        "pdf_peticao": pdf_pet_path,
        "pdf_final": pdf_final_path,
        "db": caminho_db
    }


# ============================================================
# INTERFACE TKINTER (GUI)
# ============================================================

def lancar_gui():
    """Abre interface gráfica para seleção de arquivos e processamento."""
    import tkinter as tk
    from tkinter import filedialog, messagebox, scrolledtext

    janela = tk.Tk()
    janela.title("Sistema TCE – Extração de Certidões de Débito")
    janela.configure(bg="#1A3A5C")
    janela.resizable(True, True)

    # ── Variáveis ────────────────────────────────────────────────────────────
    var_certidao  = tk.StringVar()
    var_planilha  = tk.StringVar()
    var_enderecos = tk.StringVar()   # PDF de endereços das partes
    var_pasta     = tk.StringVar(value=os.path.join(os.getcwd(), "saida_peticoes"))
    var_db        = tk.StringVar(value=os.path.join(os.getcwd(), "certidoes_tce.db"))

    # ── Helpers de layout ────────────────────────────────────────────────────
    def label(parent, text, row, col, **kw):
        tk.Label(parent, text=text, bg="#1A3A5C", fg="white",
                 font=("Helvetica", 10)).grid(row=row, column=col,
                 sticky="w", padx=8, pady=4, **kw)

    def entry(parent, textvariable, row, col, width=55):
        e = tk.Entry(parent, textvariable=textvariable, width=width,
                     font=("Helvetica", 9))
        e.grid(row=row, column=col, padx=4, pady=4, sticky="ew")
        return e

    def btn_selecionar(parent, row, col, varname, tipos):
        def sel():
            path = filedialog.askopenfilename(filetypes=tipos)
            if path:
                varname.set(path)
        tk.Button(parent, text="Selecionar", command=sel,
                  bg="#4A7FAF", fg="white", font=("Helvetica", 9),
                  relief="flat", padx=6).grid(row=row, column=col, padx=4, pady=4)

    # ── Frame principal ──────────────────────────────────────────────────────
    fr = tk.Frame(janela, bg="#1A3A5C")
    fr.pack(fill="both", expand=True, padx=16, pady=12)

    # Título
    tk.Label(fr, text="SISTEMA TCE/PI – CERTIDÕES DE DÉBITO",
             bg="#1A3A5C", fg="#F0C040",
             font=("Helvetica", 14, "bold")).grid(
             row=0, column=0, columnspan=3, pady=10)

    label(fr, "PDF da Certidão de Débito:", 1, 0)
    entry(fr, var_certidao, 1, 1)
    btn_selecionar(fr, 1, 2, var_certidao,
                   [("PDF", "*.pdf"), ("Todos", "*.*")])

    label(fr, "PDF da Planilha de Cálculo:", 2, 0)
    entry(fr, var_planilha, 2, 1)
    btn_selecionar(fr, 2, 2, var_planilha,
                   [("PDF", "*.pdf"), ("Todos", "*.*")])

    # ── Novo campo: PDF de Endereços ─────────────────────────────────────
    label(fr, "PDF de Endereços das Partes:", 3, 0)
    entry(fr, var_enderecos, 3, 1)
    btn_selecionar(fr, 3, 2, var_enderecos,
                   [("PDF", "*.pdf"), ("Todos", "*.*")])

    label(fr, "Pasta de saída:", 4, 0)
    entry(fr, var_pasta, 4, 1)
    tk.Button(fr, text="Pasta", command=lambda: var_pasta.set(
        filedialog.askdirectory() or var_pasta.get()),
        bg="#4A7FAF", fg="white", font=("Helvetica", 9),
        relief="flat", padx=6).grid(row=4, column=2, padx=4, pady=4)

    label(fr, "Banco de dados (.db):", 5, 0)
    entry(fr, var_db, 5, 1)

    # Log
    tk.Label(fr, text="Log:", bg="#1A3A5C", fg="white",
             font=("Helvetica", 10)).grid(row=6, column=0, sticky="nw",
             padx=8, pady=4)
    log_box = scrolledtext.ScrolledText(fr, width=70, height=14,
                                         font=("Courier", 9),
                                         state="disabled")
    log_box.grid(row=6, column=1, columnspan=2, padx=4, pady=4, sticky="ew")

    def log(msg):
        log_box.configure(state="normal")
        log_box.insert("end", msg + "\n")
        log_box.see("end")
        log_box.configure(state="disabled")
        janela.update()

    # ── Ação principal ───────────────────────────────────────────────────────
    def processar():
        certidao  = var_certidao.get().strip()
        planilha  = var_planilha.get().strip() or None
        enderecos = var_enderecos.get().strip() or None
        pasta     = var_pasta.get().strip()
        db        = var_db.get().strip()

        if not certidao:
            messagebox.showerror("Erro", "Selecione o PDF da certidão.")
            return

        log("=" * 60)
        log(f"Processando: {os.path.basename(certidao)}")

        try:
            import io, sys
            # Redireciona stdout para o log
            old_stdout = sys.stdout
            sys.stdout = buffer = io.StringIO()

            # enderecos é passado para processar_certidao, que importa
            # os endereços ANTES de gerar a petição (tabela já preenchida)
            resultado = processar_certidao(certidao, planilha, pasta, db,
                                           caminho_enderecos=enderecos)

            sys.stdout = old_stdout
            for linha in buffer.getvalue().splitlines():
                log(linha)

            log("\n✅ CONCLUÍDO!")
            log(f"  DOCX:      {resultado['docx']}")
            log(f"  PDF final: {resultado.get('pdf_final','N/A')}")
            log(f"  Banco DB:  {resultado['db']}")

            messagebox.showinfo("Sucesso",
                f"Processamento concluído!\n\n"
                f"DOCX: {resultado['docx']}\n"
                f"PDF:  {resultado.get('pdf_final','(conversão indisponível)')}")

        except Exception as e:
            import sys; sys.stdout = old_stdout if 'old_stdout' in dir() else sys.stdout
            log(f"❌ ERRO: {e}")
            import traceback; log(traceback.format_exc())
            messagebox.showerror("Erro", str(e))

    # ── Botão listar banco ───────────────────────────────────────────────────
    def listar_db():
        db = var_db.get().strip()
        if not os.path.exists(db):
            messagebox.showinfo("Info", "Banco ainda não criado.")
            return
        conn = inicializar_banco(db)
        certs = listar_certidoes(conn)
        conn.close()
        log("=" * 60)
        log(f"Certidões no banco: {len(certs)}")
        for c in certs:
            log(f"  [{c['id']}] {c['numero_processo']} | {c['valor_atualizado']}")
            for r in c["responsaveis"]:
                status = " [excluído]" if r["excluido"] else ""
                end_txt = f" | {r['endereco']}" if r.get("endereco") else ""
                log(f"      • {r['nome']} ({r['tipo_doc']}: {r['numero_doc']}){status}{end_txt}")
    
    # ── Botão enviar peticao ───────────────────────────────────────────────────
    def enviar_peticao():
        abrir_enviador()

    # ── Botões de ação ───────────────────────────────────────────────────────
    # ── Botões de ação ───────────────────────────────────────────────────────
    fr_btns = tk.Frame(fr, bg="#1A3A5C")
    fr_btns.grid(row=7, column=0, columnspan=3, pady=12)

    for (txt, cmd, cor) in [
        ("⚙️  Processar Certidão", processar,         "#2E7D32"),
        ("📤 Enviar Petição",      enviar_peticao, "#1565C0"),  # ✅ NOVO BOTÃO
        ("📋 Listar Banco",        listar_db,         "#4A7FAF"),
        ("❌ Fechar",              janela.quit,       "#B71C1C"),
    ]:
        tk.Button(
            fr_btns,
            text=txt,
            command=cmd,
            bg=cor,
            fg="white",
            font=("Helvetica", 11, "bold"),
            relief="flat",
            padx=14,
            pady=8
        ).pack(side="left", padx=10)

    janela.mainloop()


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) == 1:
        # Sem argumentos → GUI
        lancar_gui()
    elif len(sys.argv) >= 2:
        # Modo linha de comando: python extrator_certidao.py certidao.pdf [planilha.pdf]
        certidao_pdf  = sys.argv[1]
        planilha_pdf  = sys.argv[2] if len(sys.argv) >= 3 else None
        enderecos_pdf = sys.argv[3] if len(sys.argv) >= 4 else None
        resultado = processar_certidao(certidao_pdf, planilha_pdf,
                                       caminho_enderecos=enderecos_pdf)
        print(json.dumps({k: v for k, v in resultado.items() if k != "dados"},
                          ensure_ascii=False, indent=2))
